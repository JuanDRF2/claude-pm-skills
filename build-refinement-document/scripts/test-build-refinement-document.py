#!/usr/bin/env python3
"""Regression test for build_refinement_document.py: verify the English and
Spanish parsing paths both extract real content instead of crashing or
silently producing an empty/placeholder document."""

import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("Missing dependency: python-docx. Install it before running this test.") from exc

SCRIPT = Path(__file__).resolve().parent / "build_refinement_document.py"

EN_STORIES = """# User Stories

- Proyecto: Demo Membership

## US-DEMO-01 — Purchase individual membership

#### User story
- As a visitor
- I want to purchase an individual membership
- so that I can access member benefits

#### Scope
- Included: online checkout for individual plan
- Excluded: family plan bundling

#### Acceptance criteria

### AC-DEMO-01-01 — Membership is activated after payment

**Rules:** BR-01

**Acceptance condition:** The membership must be marked active immediately after a successful payment.

###### SC-DEMO-01-01-01 — Visitor completes checkout successfully

**Given:** the visitor has a valid payment method
**When:** the visitor submits the checkout form
**Then:**
  - the membership is created in Active status

- **Automation:** Automate now
- **Automation rationale:** Core purchase flow, high regression value
- **Automation priority:** High
- **Recommended automation level:** API
- **Automation dependencies:** None
- **Automated coverage:** Not started
"""

EN_RULES = """# Rules and Questions

| BR | Description |
|---|---|
| BR-01 | A membership is Active only after a successful payment is confirmed |
"""

EN_CHECKS = """# Test Coverage

| Check | Story | Criterion | Rules | What must be proven | Risk | Level | Evidence | Status |
|---|---|---|---|---|---|---|---|---|
| CHK-DEMO-001 | US-DEMO-01 | AC-DEMO-01-01 | BR-01 | Membership status becomes Active after payment | High | API | Membership record | Not started |
"""

EN_TESTS = """# Functional Test Cases

## FTC-DEMO-01 — Individual membership purchase

### SC-DEMO-01-01-01 — Visitor completes checkout successfully

**Covered checks:** CHK-DEMO-001

**Given:** the visitor has a valid payment method
**When:** the visitor submits the checkout form
**Then:**
  - the membership is created in Active status

- **Automation:** Automate now
- **Automation rationale:** Core purchase flow, high regression value
- **Automation priority:** High
- **Recommended automation level:** API
- **Automation dependencies:** None
- **Automated coverage:** Not started
"""

ES_STORIES = EN_STORIES.replace(
    "#### User story", "#### Historia"
).replace(
    "- As a visitor\n- I want to purchase an individual membership\n- so that I can access member benefits",
    "- Como visitante\n- Quiero comprar una membresía individual\n- para poder acceder a los beneficios de miembro",
).replace(
    "#### Scope", "#### Alcance y dependencias"
).replace(
    "- Included: online checkout for individual plan\n- Excluded: family plan bundling",
    "- Incluye: checkout en línea para el plan individual\n- Excluye: combinación con el plan familiar",
).replace(
    "**Rules:**", "**Reglas:**"
).replace(
    "**Given:**", "**Dado:**"
).replace(
    "**When:**", "**Cuando:**"
).replace(
    "**Then:**", "**Entonces:**"
).replace(
    "- **Automation:**", "- **Automatización:**"
).replace(
    "- **Automation rationale:**", "- **Razón de automatización:**"
).replace(
    "- **Automation priority:**", "- **Prioridad de automatización:**"
).replace(
    "- **Recommended automation level:**", "- **Nivel de automatización recomendado:**"
).replace(
    "- **Automation dependencies:**", "- **Dependencias de automatización:**"
).replace(
    "- **Automated coverage:**", "- **Cobertura automatizada:**"
)

ES_CHECKS = EN_CHECKS.replace("What must be proven", "Qué debe comprobarse")
ES_TESTS = EN_TESTS.replace(
    "**Given:**", "**Dado:**"
).replace(
    "**When:**", "**Cuando:**"
).replace(
    "**Then:**", "**Entonces:**"
).replace(
    "- **Automation:**", "- **Automatización:**"
).replace(
    "- **Automation rationale:**", "- **Razón de automatización:**"
).replace(
    "- **Automation priority:**", "- **Prioridad de automatización:**"
).replace(
    "- **Recommended automation level:**", "- **Nivel de automatización recomendado:**"
).replace(
    "- **Automation dependencies:**", "- **Dependencias de automatización:**"
).replace(
    "- **Automated coverage:**", "- **Cobertura automatizada:**"
)


def write_fixture(root: Path, stories: str, rules: str, checks: str, tests: str) -> None:
    root.mkdir(parents=True, exist_ok=True)
    (root / "05-user-stories.md").write_text(stories, encoding="utf-8")
    (root / "02-rules-and-questions.md").write_text(rules, encoding="utf-8")
    (root / "06-test-coverage.md").write_text(checks, encoding="utf-8")
    (root / "07-functional-test-cases.md").write_text(tests, encoding="utf-8")
    (root / "08-traceability-and-risks.md").write_text("", encoding="utf-8")


def build(root: Path, output: Path) -> subprocess.CompletedProcess:
    return subprocess.run(
        [sys.executable, str(SCRIPT), str(root), "--output", str(output)],
        capture_output=True,
        text=True,
    )


def doc_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def main() -> int:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)

        en_root = tmp_path / "en-package"
        write_fixture(en_root, EN_STORIES, EN_RULES, EN_CHECKS, EN_TESTS)
        en_output = tmp_path / "en.docx"
        en_result = build(en_root, en_output)
        assert en_result.returncode == 0, (
            f"English package failed to build:\nstdout={en_result.stdout}\nstderr={en_result.stderr}"
        )
        assert en_output.exists(), "English .docx was not created"
        en_text = doc_text(en_output)
        assert "the visitor has a valid payment method" in en_text, "English Given value missing from document"
        assert "the membership is created in Active status" in en_text, "English Then value missing from document"
        assert "A membership is Active only after a successful payment is confirmed" in en_text, (
            "English BR-01 rule text missing from document"
        )
        assert "Membership status becomes Active after payment" in en_text, (
            "English check description ('What must be proven' column) missing from document"
        )
        print("PASS: English package builds and extracts Given/When/Then, rules, and check fields")

        es_root = tmp_path / "es-package"
        write_fixture(es_root, ES_STORIES, EN_RULES, ES_CHECKS, ES_TESTS)
        es_output = tmp_path / "es.docx"
        es_result = build(es_root, es_output)
        assert es_result.returncode == 0, (
            f"Spanish package failed to build (regression):\nstdout={es_result.stdout}\nstderr={es_result.stderr}"
        )
        assert es_output.exists(), "Spanish .docx was not created"
        es_text = doc_text(es_output)
        assert "the visitor has a valid payment method" in es_text, "Spanish Dado value missing from document"
        assert "the membership is created in Active status" in es_text, "Spanish Entonces value missing from document"
        print("PASS: Spanish package still builds and extracts Dado/Cuando/Entonces fields (no regression)")

    print("ALL BUILD-REFINEMENT-DOCUMENT CHECKS PASSED")
    return 0


if __name__ == "__main__":
    sys.exit(main())
