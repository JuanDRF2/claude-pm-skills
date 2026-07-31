#!/usr/bin/env python3
"""Build a readable DOCX view from the canonical refinement Markdown package."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:
    raise SystemExit("Missing dependency: python-docx. Install it in an approved Python environment.") from exc


BLUE = "1F4E79"
LIGHT_BLUE = "EAF2F8"
GREEN = "287A5A"
LIGHT_GREEN = "EAF6F0"
INK = "172033"
MUTED = "5F6B76"


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def split_pipe(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def metadata(text: str) -> dict[str, str]:
    out = {}
    for line in text.splitlines():
        match = re.match(r"^-\s+([^:]+):\s*(.+)$", line)
        if match:
            out[match.group(1).strip()] = match.group(2).strip()
    return out


def parse_rules(text: str) -> dict[str, str]:
    rules = {}
    for line in text.splitlines():
        cells = split_pipe(line) if line.startswith("|") else []
        if len(cells) >= 2 and re.fullmatch(r"BR-[A-Z0-9-]+", cells[0]):
            rules[cells[0]] = cells[1]
    return rules


def expand_rule_references(value: str) -> list[str]:
    """Expand BR namespace ranges while preserving first-seen order."""
    result: list[str] = []
    token_pattern = re.compile(
        r"\b(BR-(?:(?P<namespace>[A-Z0-9]+)-)?(?P<start>\d{2,}))"
        r"(?:\s*[–—-]\s*(?:(?:BR-)?(?:(?P<end_namespace>[A-Z0-9]+)-)?(?P<end>\d{2,})))?"
    )
    for match in token_pattern.finditer(value):
        prefix = f"BR-{match.group('namespace')}-" if match.group("namespace") else "BR-"
        start = int(match.group("start"))
        end = int(match.group("end") or start)
        width = max(len(match.group("start")), len(match.group("end") or ""))
        if end < start or end - start > 500:
            continue
        for number in range(start, end + 1):
            rule_id = f"{prefix}{number:0{width}d}"
            if rule_id not in result:
                result.append(rule_id)
    return result


def parse_checks(text: str) -> dict[str, dict[str, str]]:
    checks = {}
    headers = []
    for line in text.splitlines():
        if not line.startswith("|"):
            continue
        cells = split_pipe(line)
        if cells and cells[0].lower() == "check":
            headers = cells
        elif headers and cells and re.fullmatch(r"CHK-[A-Z0-9-]+", cells[0]):
            checks[cells[0]] = dict(zip(headers, cells))
    return checks


def blocks(text: str, pattern: str) -> list[tuple[str, str, str]]:
    matches = list(re.finditer(pattern, text, re.M))
    result = []
    for idx, match in enumerate(matches):
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        result.append((match.group(1), match.group(2).strip(), text[match.end():end].strip()))
    return result


SCENARIO_FIELD_ALIASES = {
    "dado": "dado", "given": "dado",
    "cuando": "cuando", "when": "cuando",
    "entonces": "entonces", "then": "entonces",
    "automatización": "automatización", "automation": "automatización",
    "razón de automatización": "razón de automatización", "automation rationale": "razón de automatización",
    "prioridad de automatización": "prioridad de automatización", "automation priority": "prioridad de automatización",
    "nivel de automatización recomendado": "nivel de automatización recomendado",
    "recommended automation level": "nivel de automatización recomendado",
    "dependencias de automatización": "dependencias de automatización",
    "automation dependencies": "dependencias de automatización",
    "cobertura automatizada": "cobertura automatizada", "automated coverage": "cobertura automatizada",
}
SCENARIO_FIELD_PATTERN = "|".join(re.escape(label) for label in SCENARIO_FIELD_ALIASES)


def parse_scenarios(text: str) -> dict[str, dict[str, str]]:
    result = {}
    matches = list(re.finditer(r"^###\s+(SC-[A-Z0-9-]+)\s+[—-]\s+(.+)$", text, re.M))
    for idx, match in enumerate(matches):
        next_case = re.search(r"^##\s+FTC-[A-Z0-9-]+\s+[—-]", text[match.end():], re.M)
        next_scenario = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        next_case_at = match.end() + next_case.start() if next_case else len(text)
        end = min(next_scenario, next_case_at)
        code, title, body = match.group(1), match.group(2).strip(), text[match.end():end].strip()
        fields = {"title": title, "body": body}
        lines = body.splitlines()
        for line_idx, line in enumerate(lines):
            field = re.match(rf"^\s*(?:-\s*)?\*\*({SCENARIO_FIELD_PATTERN}):\*\*\s*(.*)$", line, re.I)
            if not field:
                continue
            label = SCENARIO_FIELD_ALIASES[field.group(1).lower()]
            value = field.group(2).strip()
            if label == "entonces" and not value:
                following = []
                for extra in lines[line_idx + 1:]:
                    if not extra.strip() or re.match(r"^\s*(?:-\s*)?\*\*", extra):
                        break
                    following.append(re.sub(r"^\s*-\s*", "", extra).strip())
                value = " ".join(following)
            fields[label] = re.sub(r"\s+", " ", value).strip(" -\n")
        result[code] = fields
    return result


def parse_functional_cases(text: str) -> dict[str, dict[str, object]]:
    result = {}
    for code, title, body in blocks(text, r"^##\s+(FTC-[A-Z0-9-]+)\s+[—-]\s+(.+)$"):
        result[code] = {
            "title": title,
            "body": body,
            "scenarios": re.findall(r"^###\s+(SC-[A-Z0-9-]+)\s+[—-]\s+(.+)$", body, re.M),
        }
    return result


def section(body: str, heading: str) -> str:
    match = re.search(rf"^###\s+{re.escape(heading)}\s*$\n(.*?)(?=^###\s+|\Z)", body, re.M | re.S)
    return match.group(1).strip() if match else ""


def story_section(body: str, heading: str) -> str:
    match = re.search(
        rf"^\*\*{re.escape(heading)}\*\*\s*$\n(.*?)(?=^\*\*[^*\n]+\*\*\s*$|^###\s+|\Z)",
        body,
        re.M | re.S,
    )
    return match.group(1).strip() if match else section(body, heading)


def parse_criteria(story_body: str) -> list[dict[str, object]]:
    area = section(story_body, "Criterios de aceptación") or section(story_body, "Acceptance criteria")
    if not area:
        area = story_body
    story_rule_match = re.search(r"\*\*(?:Reglas|Rules):\*\*\s*([^\n]+)", story_body)
    story_rules = expand_rule_references(story_rule_match.group(1)) if story_rule_match else []
    result = []
    criterion_pattern = r"^#{3,4}\s+(AC-[A-Z0-9-]+)\s+[—-]\s+(.+)$"
    for code, title, body in blocks(area, criterion_pattern):
        rule_match = re.search(r"\*\*(?:Reglas|Rules):\*\*\s*([^\n]+)", body)
        rules = expand_rule_references(rule_match.group(1)) if rule_match else story_rules
        condition_match = re.search(r"\*\*(?:Condición de aceptación|Acceptance condition):\*\*\s*([^\n]+)", body)
        gherkin = []
        for line in body.splitlines():
            match = re.match(r"\*\*(Dado que|Dado|Given that|Given|Y|And|Cuando|When|Entonces|Then|Pero|But)\*\*\s*(.+?)\s*$", line, re.I)
            if match:
                gherkin.append((match.group(1), match.group(2)))
        result.append({"code": code, "title": title, "rules": rules, "condition": condition_match.group(1).strip() if condition_match else "", "gherkin": gherkin})
    return result


def shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def font(run, size=10.5, bold=False, color=INK) -> None:
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, before, after in (("Title", 25, 0, 10), ("Heading 1", 17, 16, 8), ("Heading 2", 13.5, 12, 6), ("Heading 3", 11.5, 8, 4)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def labeled(doc: Document, label: str, value: str) -> None:
    if not value:
        return
    p = doc.add_paragraph()
    font(p.add_run(label + " "), bold=True, color=BLUE)
    font(p.add_run(re.sub(r"\s+", " ", value).strip()))


def callout(doc: Document, label: str, value: str, green=False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GREEN if green else LIGHT_BLUE)
    p = cell.paragraphs[0]
    font(p.add_run(label + "\n"), bold=True, color=GREEN if green else BLUE)
    font(p.add_run(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_gherkin(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    for idx, (key, value) in enumerate(rows):
        shade(table.cell(idx, 0), "F4F7FA")
        p = table.cell(idx, 0).paragraphs[0]
        font(p.add_run(key), bold=True, color=BLUE)
        p = table.cell(idx, 1).paragraphs[0]
        font(p.add_run(value))


def related_checks(criterion: str, checks: dict[str, dict[str, str]]) -> list[dict[str, str]]:
    return [row for row in checks.values() if criterion in (row.get("Criterio") or row.get("Criterion") or "")]


def row_value(row: dict[str, str], *names: str) -> str:
    return next((row[name] for name in names if row.get(name)), "")


def add_story(doc: Document, story, rules, checks, scenarios, cases) -> None:
    code, title, body = story
    doc.add_page_break()
    p = doc.add_paragraph()
    font(p.add_run(f"Historia de usuario · {code}"), 9.5, True, BLUE)
    doc.add_paragraph(title, style="Title")

    history = (
        story_section(body, "Historia")
        or story_section(body, "Historia de usuario")
        or story_section(body, "User story")
        or story_section(body, "Use Case")
    )
    history_text = " ".join(
        re.sub(r"^-\s*", "", line).strip()
        for line in history.splitlines()
        if line.strip()
    )
    callout(doc, "Historia", history_text or "Descripción no disponible en la fuente.")

    doc.add_paragraph("Alcance y decisiones", style="Heading 1")
    scope = (
        story_section(body, "Alcance y dependencias")
        or story_section(body, "Alcance")
        or story_section(body, "Scope and dependencies")
        or story_section(body, "Scope")
    )
    for line in scope.splitlines():
        if line.startswith("-"):
            value = line[1:].strip()
            if ":" in value:
                key, rest = value.split(":", 1)
                labeled(doc, key + ":", rest)
            else:
                labeled(doc, "Detalle:", value)

    doc.add_paragraph("Criterios de aceptación y cobertura", style="Heading 1")
    for criterion in parse_criteria(body):
        doc.add_paragraph(f"Criterio de aceptación · {criterion['code']}", style="Heading 2")
        p = doc.add_paragraph()
        font(p.add_run(str(criterion["title"])), 12, True)
        labeled(doc, "Condición de aceptación:", str(criterion.get("condition", "")))
        if criterion["gherkin"]:
            add_gherkin(doc, criterion["gherkin"])
        doc.add_paragraph("Reglas de negocio aplicables", style="Heading 3")
        for rule_id in criterion["rules"]:
            labeled(doc, f"Regla de negocio · {rule_id}:", rules.get(rule_id, "Definición no disponible."))
        doc.add_paragraph("Cómo se comprobará", style="Heading 3")
        for check in related_checks(str(criterion["code"]), checks):
            callout(doc, f"Comprobación de cobertura · {check.get('Check', '')}", row_value(check, "Qué debe comprobarse", "Qué debe probarse", "What must be proven"), green=True)
            labeled(doc, "Riesgo:", row_value(check, "Riesgo", "Risk"))
            labeled(doc, "Nivel recomendado:", row_value(check, "Nivel", "Level"))
            labeled(doc, "Evidencia esperada:", row_value(check, "Evidencia", "Evidence"))
            check_id = check.get("Check", "")
            for sc_id, sc in scenarios.items():
                if check_id and check_id in sc.get("body", ""):
                    doc.add_paragraph(f"Escenario de prueba · {sc_id}", style="Heading 3")
                    labeled(doc, "Título:", sc.get("title", ""))
                    labeled(doc, "Dado:", sc.get("dado", ""))
                    labeled(doc, "Cuando:", sc.get("cuando", ""))
                    labeled(doc, "Entonces:", sc.get("entonces", ""))
                    labeled(doc, "Automatización:", sc.get("automatización", ""))
                    labeled(doc, "Nivel recomendado:", sc.get("nivel de automatización recomendado", ""))
                    labeled(doc, "Prioridad:", sc.get("prioridad de automatización", ""))
                    labeled(doc, "Razón:", sc.get("razón de automatización", ""))
                    labeled(doc, "Dependencias:", sc.get("dependencias de automatización", ""))
                    labeled(doc, "Estado:", sc.get("cobertura automatizada", ""))

    related_case_ids = set(re.findall(r"FTC-[A-Z0-9-]+", body))
    if related_case_ids:
        doc.add_paragraph("Casos funcionales relacionados", style="Heading 1")
        for case_id in sorted(related_case_ids):
            case = cases.get(case_id)
            if not case:
                continue
            doc.add_paragraph(f"Caso funcional · {case_id} — {case['title']}", style="Heading 2")
            for sc_id, sc_title in case["scenarios"]:
                labeled(doc, f"Escenario de prueba · {sc_id}:", sc_title)


def build(artifact_dir: Path, output: Path) -> None:
    stories_text = read(artifact_dir / "05-user-stories.md")
    if not stories_text:
        raise SystemExit("Missing required source: 05-user-stories.md")
    stories = blocks(stories_text, r"^##\s+(US-[A-Z0-9-]+)\s+[—-]\s+(.+)$")
    if not stories:
        raise SystemExit("No US-* stories found in 05-user-stories.md")

    project = metadata(stories_text).get("Proyecto", artifact_dir.name.replace("-", " ").title())
    rules = parse_rules(read(artifact_dir / "02-rules-and-questions.md"))
    checks = parse_checks(read(artifact_dir / "06-test-coverage.md"))
    tests_text = read(artifact_dir / "07-functional-test-cases.md")
    scenarios = parse_scenarios(tests_text)
    cases = parse_functional_cases(tests_text)
    strategy_fields = {
        "automatización",
        "nivel de automatización recomendado",
        "prioridad de automatización",
        "razón de automatización",
        "dependencias de automatización",
        "cobertura automatizada",
    }
    incomplete = {code: sorted(strategy_fields - set(data)) for code, data in scenarios.items() if strategy_fields - set(data)}
    if incomplete:
        details = "; ".join(f"{code}: {', '.join(fields)}" for code, fields in incomplete.items())
        raise SystemExit(f"Incomplete canonical QA strategy; Word publication stopped. {details}")

    doc = Document()
    configure(doc)
    doc.add_paragraph("Documento de refinamiento", style="Heading 3")
    doc.add_paragraph(project, style="Title")
    callout(doc, "Propósito", "Reunir historias, criterios, reglas y pruebas relacionadas en una lectura documental clara.")
    judge_report = read(artifact_dir / "11-refinement-judge-report.md")
    if judge_report:
        verdict_match = re.search(
            r"^-\s*(?:Verdict|Veredicto)(?:\s*/\s*(?:Verdict|Veredicto))?\s*:\s*(.+)$",
            judge_report,
            re.M | re.I,
        )
        gate_match = re.search(
            r"^-\s*(?:Gate decision|Decisión del gate)(?:\s*/\s*(?:Gate decision|Decisión del gate))?\s*:\s*(.+)$",
            judge_report,
            re.M | re.I,
        )
        callout(
            doc,
            "Refinement Judge",
            f"Veredicto: {verdict_match.group(1).strip() if verdict_match else 'No reconocido'}. "
            f"Decisión del gate: {gate_match.group(1).strip() if gate_match else 'No disponible'}.",
        )
    doc.add_paragraph("Cómo leer", style="Heading 1")
    for label, meaning in (
        ("Historia de usuario (US):", "necesidad y valor."),
        ("Criterio de aceptación (AC):", "condición observable."),
        ("Regla de negocio (BR):", "comportamiento acordado."),
        ("Comprobación de cobertura (CHK):", "aspecto concreto que QA debe demostrar."),
        ("Caso funcional (FTC):", "grupo coherente de escenarios."),
        ("Escenario de prueba (SC):", "recorrido ejecutable."),
    ):
        labeled(doc, label, meaning)
    labeled(doc, "Contenido:", f"{len(stories)} historias, {sum(len(parse_criteria(body)) for _, _, body in stories)} criterios, {len(checks)} checks, {len(cases)} casos funcionales y {len(scenarios)} escenarios detectados.")

    for story in stories:
        add_story(doc, story, rules, checks, scenarios, cases)

    doc.add_page_break()
    doc.add_paragraph("Pendientes y riesgos", style="Title")
    risks = read(artifact_dir / "08-traceability-and-risks.md")
    for heading in ("Elementos bloqueados", "Revisiones pendientes", "Riesgo residual"):
        match = re.search(rf"^##\s+{re.escape(heading)}\s*$\n(.*?)(?=^##\s+|\Z)", risks, re.M | re.S)
        if match:
            doc.add_paragraph(heading, style="Heading 1")
            for line in match.group(1).splitlines():
                if line.strip() and not line.startswith("|") and not re.match(r"^[-| ]+$", line):
                    doc.add_paragraph(re.sub(r"^[-# ]+", "", line).strip())

    if judge_report:
        doc.add_page_break()
        doc.add_paragraph("Refinement Judge", style="Title")
        for line in judge_report.splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith("# Refinement Judge"):
                continue
            heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
            if heading:
                level = min(len(heading.group(1)) - 1, 3)
                doc.add_paragraph(heading.group(2), style=f"Heading {level}")
            elif stripped.startswith(("- ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif not stripped.startswith("|"):
                doc.add_paragraph(stripped)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = f"Documento de refinamiento — {project}"
    doc.core_properties.author = "Product Refinement Skills"
    doc.save(output)
    print(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("artifact_dir", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.artifact_dir.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
